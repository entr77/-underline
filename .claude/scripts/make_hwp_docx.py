"""
서식 2 구조로 사업계획서 DOCX 생성 v2 — AI 전략 + RAG + 수익화 업데이트
다이어그램은 텍스트 박스 스타일 표로 구현 (HWP 호환성 최우선)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── 기본 스타일 ───────────────────────────────
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)

for section in doc.sections:
    section.page_width   = Cm(21)
    section.page_height  = Cm(29.7)
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin   = Cm(2.5)
    section.bottom_margin= Cm(2.5)

# ─── 헬퍼 함수 ────────────────────────────────
FOREST = RGBColor(0x1E, 0x3A, 0x2F)
INK    = RGBColor(0x1C, 0x19, 0x17)
GRAY   = RGBColor(0x80, 0x80, 0x80)
HDR_FILL = 'D9E8DC'   # 헤더 연두
SUB_FILL = 'F2F7F3'   # 서브헤더 연한 녹색

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = FOREST
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = INK
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)

def add_body(text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run('※ ' + text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    p.paragraph_format.space_after = Pt(4)

def sp(n=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(n)

def add_table(headers, rows, col_widths=None, sub_header_rows=None):
    """sub_header_rows: set of row indices (0-based in rows) that should get sub-fill"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        shade_cell(hdr_cells[i], HDR_FILL)
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row):
            row_cells[c_idx].text = str(cell_text)
            for run in row_cells[c_idx].paragraphs[0].runs:
                run.font.size = Pt(9)
        if sub_header_rows and r_idx in sub_header_rows:
            for cell in row_cells:
                shade_cell(cell, SUB_FILL)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[i])
    sp(3)
    return table

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C0C0C0')
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)

def add_diagram_box(lines, title=None):
    """텍스트 박스 스타일 다이어그램 — 1열 표로 구현"""
    rows_data = []
    if title:
        rows_data.append([title])
    for line in lines:
        rows_data.append([line])
    table = doc.add_table(rows=len(rows_data), cols=1)
    table.style = 'Table Grid'
    for i, row_data in enumerate(rows_data):
        cell = table.rows[i].cells[0]
        cell.text = row_data[0]
        p = cell.paragraphs[0]
        for run in p.runs:
            run.font.name = '맑은 고딕'
            run.font.size = Pt(9)
            if title and i == 0:
                run.bold = True
                run.font.color.rgb = FOREST
                shade_cell(cell, HDR_FILL)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0].cells[0].width = Cm(15.5)
    sp(3)
    return table

def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('2026년 혁신 소상공인 AI 활용지원 사업\n사업계획서')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = FOREST
title_p.paragraph_format.space_after = Pt(4)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('신청유형: 사업 고도화  ｜  활용 AI: Claude  ｜  신청권역: 수도권').font.size = Pt(10)
sub.paragraph_format.space_after = Pt(2)

dt = doc.add_paragraph()
dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
dt.add_run('2026년 6월  ｜  (주)얼리커뮤니케이션').font.size = Pt(10)
dt.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
p.add_run('■ 사업아이템 한줄요약').bold = True

summary = doc.add_paragraph()
summary.add_run(
    '독자들이 책을 읽다 멈춘 문장을 Claude AI로 기록하고, '
    '같은 문장에서 멈춘 독자들을 연결하는 소셜 독서 플랫폼 고도화'
).font.size = Pt(10)
summary.paragraph_format.space_after = Pt(10)

add_divider()

# ═══════════════════════════════════════════════════════════
# PAGE 1 — 일반현황 ①
# ═══════════════════════════════════════════════════════════
add_h1('[ 일반 현황 ]')

add_h2('① 기업 소개')
add_body(
    '(주)얼리커뮤니케이션은 종이책 독자들의 밑줄 경험을 디지털화하는 소셜 독서 플랫폼 '
    '밑줄(Underline)을 운영합니다. 독자가 책을 읽다 형광펜·볼펜으로 표시한 문장을 '
    '사진 한 장으로 촬영하면, Claude Vision AI가 밑줄 친 문장과 책 정보를 자동 인식해 저장합니다. '
    '저장된 문장은 소셜 피드로 공개되어 같은 문장에서 멈춘 독자들과 연결됩니다.'
)
add_note('편집자가 고른 명언이 아닌, 실제 독자들이 읽다 멈춘 문장들의 집합 — 누가 멈췄는지가 그 문장의 맥락이 됩니다.')
add_table(
    ['항목', '내용'],
    [
        ['서비스명', '밑줄 (Underline)'],
        ['핵심 AI', 'Claude Vision API (claude-sonnet-4-6)'],
        ['운영 현황', 'MVP 프로덕션 운영 중  ｜  누적 밑줄 6,227건  ｜  등록 도서 264권'],
        ['기술 스택', 'Next.js 15, Supabase PostgreSQL + pgvector, Vercel'],
    ],
    col_widths=[3.5, 12]
)

add_h2('② 지원 동기')
add_body('이미 검증된 Claude AI를 고도화하고, 자체 RAG 엔진을 구축합니다.', bold=True)
add_body(
    'Claude Vision API를 프로덕션에서 운영 중이며 6,227건의 실사용 데이터로 핵심 파이프라인이 검증되었습니다. '
    '이번 사업을 통해 ① AI 자동 분석(장르·성향·클러스터링) ② 자체 RAG 구축 ③ 개인화 추천·수익화로 확장합니다.'
)
add_table(
    ['목표', '내용', '기대효과'],
    [
        ['Phase 1', 'AI 자동 분석 — 장르·태그·유저 성향·동일 밑줄 클러스터링', '수동 입력 70% 단축'],
        ['Phase 2', '자체 RAG 구축 — 취향 피드 개인화 + 책 요약표 자동 생성', 'Claude 비용 30%↓'],
        ['Phase 3', '추천 엔진 — AI 책 추천 + 상황별 인용문 추천 + 수익화 연결', '첫 수익 실현'],
    ],
    col_widths=[2.5, 9, 4]
)

add_h2('③ 조직 구성')
add_table(
    ['구분', '역할', '담당 업무'],
    [
        ['대표 (1인)', '개발·기획·운영', 'Claude API 연동, Supabase 운영, 서비스 개발, 마케팅'],
        ['채용 예정 (사업 완료 후)', '마케팅 담당 1인', '저자 파트너십, 북클럽 이벤트 운영'],
    ],
    col_widths=[3.5, 3, 9]
)
add_note('이미지 분석·책 정보 매칭·취향 집계·배포 자동화 완료. Claude AI 자동화로 1인 운영 가능한 구조 구축.')

page_break()

# ═══════════════════════════════════════════════════════════
# PAGE 2 — 일반현황 ②
# ═══════════════════════════════════════════════════════════

add_h2('④ 기업 현황')
add_table(
    ['항목', '현황'],
    [
        ['현재 매출', '0원 (MVP 검증 단계 — 수익화 기능 미구현 상태)'],
        ['실사용 데이터', '누적 밑줄 6,227건  ｜  등록 도서 264권  ｜  가입자 약 100명'],
        ['AI 기술 지표', 'Claude Vision OCR 정확도 95%+  ｜  한국어 종이책 특화 파이프라인'],
        ['지식재산권', '출원 준비 중'],
    ],
    col_widths=[4, 11.5]
)
add_note('매출 0원은 수익화 기능 미구현 때문. 본 사업 완료 후 6개월 내 프리미엄 구독 출시로 수익 시작.')

add_divider()
add_h2('⑤ 사업 내용')

add_body('【시장성】 독자들의 밑줄 공유 행동은 이미 대규모로 존재합니다.', bold=True)
add_table(
    ['지표', '수치'],
    [
        ['전 세계 #bookstagram 게시물', '1억 1,900만 건 이상 (2025년)'],
        ['BookTok 누적 조회수', '2,000억 회 이상 (2025년)'],
        ['2024 서울국제도서전 방문객', '역대 최대 15만 명'],
        ['한국 전자출판 합산 매출', '1조 5,959억 원 (2024년, +12.2%)'],
    ],
    col_widths=[6, 9.5]
)
add_note('독자들은 이미 인스타그램에 밑줄 사진을 올립니다. 그 행동을 연결할 인프라가 없을 뿐입니다.')

add_body('【경쟁력】 세계에 없는 조합', bold=True)
add_table(
    ['', '명언앱/굿리즈', 'Readwise ($9.99/월)', '밑줄 (무료)'],
    [
        ['종이책 + 한국어', '✗', '✗', '✓'],
        ['소셜 레이어', '✗', '✗', '✓'],
        ['문장 맥락 (독자 프로필)', '✗', '✗', '✓'],
        ['RAG 기반 인용문 추천', '✗', '✗', '✓ (구현 예정)'],
    ],
    col_widths=[4.5, 3, 4, 4]
)
add_note('독자가 쌓일수록 "멈춘 독자들의 프로필" 데이터가 풍부해져 데이터 가치가 함께 커지는 네트워크 효과 발생.')

add_divider()
add_h2('⑥ AI활용 현황')
add_table(
    ['AI 기능', '역할', '누적 처리량'],
    [
        ['OCR', '책 페이지 전체 텍스트 추출', '6,227건'],
        ['밑줄 감지', '형광펜·볼펜·동그라미 자동 인식', '6,227건'],
        ['책 식별', '헤더/푸터 분석 → 카카오 도서 API 매칭', '264권'],
    ],
    col_widths=[3, 9, 3.5]
)
add_note('단일 API 호출로 3가지 기능 동시 처리(비용·속도 최적화). 이미 검증된 파이프라인 위에 취향 분석·추천 레이어 추가.')

page_break()

# ═══════════════════════════════════════════════════════════
# PAGE 3 — 추진계획 ①
# ═══════════════════════════════════════════════════════════
add_h1('[ 추진 계획 ]')

add_h2('⑦ AI 활용 아이템 소개 — Claude + 자체 RAG 이중 구조')

add_body('핵심 전략: 데이터가 쌓일수록 Claude 의존도를 낮추고 자체 RAG로 전환한다', bold=True)

# AI 전략 구조 다이어그램
add_diagram_box([
    '[ 현재: 데이터 축적기 ]          [ Phase 2~: RAG 전환기 ]          [ 장기: 자립기 ]',
    '          │                                   │                               │',
    '  Claude API 전담            Claude + pgvector RAG 병행          자체 RAG 주도',
    '  OCR·분석·추천 모두          반복 쿼리 → RAG                    Claude는 보완 레이어',
    '  Claude 처리                창의적 분석 → Claude               API 비용 70%↓',
    '',
    '  기술 기반: Supabase pgvector (이미 적용 가능한 확장 모듈) — 별도 인프라 추가 불필요',
], title='AI 전략 로드맵 — Claude → 자체 RAG 단계적 전환')

add_body('AI 기능 전체 맵', bold=True)
add_table(
    ['기능', '현재', '고도화 후', 'AI 방식'],
    [
        ['책 사진 → 도서 추출', '✅ 운영 중', '유지', 'Claude Vision'],
        ['밑줄 문장 OCR', '✅ 운영 중', '유지', 'Claude Vision'],
        ['장르·태그 자동 분류', '✗ 수동', '자동화', 'Claude → RAG 전환'],
        ['유저 성향 분석 (직업·독서 패턴)', '△ 기초 집계', '심층 분석', 'Claude 분석'],
        ['동일 밑줄 클러스터링', '✗', '자동화', 'pgvector 유사도'],
        ['밑줄별 공명 독자 프로파일링', '△ 기초', '자동화', 'RAG + 통계'],
        ['책 요약표 (책별 밑줄 전체 분석)', '✗', '자동화', 'Claude → RAG'],
        ['AI 책 추천', '✗', '자동화', 'RAG + Claude'],
        ['상황별 인용문 추천', '✗', '자동화', 'RAG 검색'],
    ],
    col_widths=[4.5, 2, 2.5, 6.5],
    sub_header_rows={0, 1}
)

page_break()

# ═══════════════════════════════════════════════════════════
# PAGE 4 — 추진계획 ②  AI 구축 계획
# ═══════════════════════════════════════════════════════════

add_h2('⑧ AI활용 모델 구축 계획 (Phase별 상세)')

# Phase 흐름 다이어그램
add_diagram_box([
    '  Phase 1 (1~2개월)           Phase 2 (3~4개월)            Phase 3 (5~6개월)',
    '  ┌─────────────────┐         ┌─────────────────┐          ┌─────────────────┐',
    '  │ AI 자동 분석 엔진 │────────▶│  RAG 구축 +     │─────────▶│  추천 엔진 +    │',
    '  │                 │         │  책 요약표       │          │  수익화 연결    │',
    '  │ ① 장르·태그 분류 │         │                 │          │                 │',
    '  │ ② 유저 성향 분석 │         │ ① pgvector DB   │          │ ① AI 책 추천    │',
    '  │ ③ 밑줄 클러스터링│         │ ② 취향 피드 개인화│         │ ② 인용문 추천   │',
    '  └─────────────────┘         │ ③ 책 요약표 생성 │          │ ③ 어필리에이트  │',
    '                              └─────────────────┘          └─────────────────┘',
    '  기대효과: 태그입력 70%↓      기대효과: Claude비용 30%↓     기대효과: 첫 수익 실현',
], title='Phase별 AI 구축 로드맵')

add_table(
    ['Phase', '기간', '핵심 개발 내용', '산출물', '기대효과'],
    [
        ['Phase 1',
         '1~2개월',
         '① 밑줄 문장 → 장르·태그 자동 분류 (Claude)\n② 유저 직업·독서 성향 심층 프로파일링\n③ 동일 밑줄 pgvector 유사도 클러스터링',
         '/api/tags/suggest\n/api/user/profile-ai\n/api/underlines/cluster',
         '태그 입력 70%↓\n유저 프로파일 자동화'],
        ['Phase 2',
         '3~4개월',
         '① pgvector 자체 벡터 DB 구축 (RAG 기반)\n② 취향 기반 피드 개인화\n   ("나와 비슷한 독자의 밑줄" 우선 노출)\n③ 책별 밑줄 자동 요약표 생성\n   (저자/출판사 대시보드 기반)',
         'pgvector 인덱스\n/api/feed/personalized\n/api/book/summary-ai',
         'Claude 비용 30%↓\n체류시간 향상\n출판사 B2B 기반'],
        ['Phase 3',
         '5~6개월',
         '① AI 책 추천 (밑줄 히스토리 → 다음 책 3권)\n② 상황별 인용문 추천\n   (직업·상황 입력 → RAG 검색)\n③ 인용 서비스·어필리에이트 수익 연결',
         '/api/book/recommend\n/api/quote/suggest\n프리미엄 결제 연동',
         'CPS 수익 시작\n프리미엄 전환 유도'],
    ],
    col_widths=[1.8, 1.5, 6, 3.2, 3]
)
add_note('user_taste_profiles 테이블 이미 구축 완료. Supabase pgvector 확장 적용으로 RAG 전환 시 별도 인프라 불필요.')

page_break()

# ═══════════════════════════════════════════════════════════
# PAGE 5 — 추진계획 ③ + 성과목표
# ═══════════════════════════════════════════════════════════

add_h2('⑨ AI 비즈니스 모델 개선 계획 — 등록 개수 제한 기반 프리미엄')

add_table(
    ['기능', '무료 (월 10개 제한)', '프리미엄 (월 4,900원)'],
    [
        ['월 밑줄 등록', '10개 제한', '무제한'],
        ['소셜 피드·좋아요', '✓', '✓'],
        ['공유 카드 테마', '기본 3종', '6종 전체 + 워터마크 제거'],
        ['AI 책 요약표', '✗', '✓'],
        ['AI 상황별 인용문 추천', '✗', '✓'],
        ['AI 취향 피드·책 추천', '✗', '✓'],
        ['인용 서비스 (출처 + 독자 맥락)', '✗', '✓'],
    ],
    col_widths=[5.5, 4, 6],
    sub_header_rows={0}
)
add_note('월 10개는 2~3권 독자에게 충분 → 부담 없이 유입. 4권+·크리에이터는 자연스럽게 한도 도달 → "더 쓰고 싶어서" 결제.')

add_body('수익 레이어 및 전환율', bold=True)
add_table(
    ['시점', 'MAU', '전환율', '유료', '월 수익', '주요 수익원'],
    [
        ['6개월', '1,000명', '5%', '50명', '245,000원', '프리미엄 + 어필리에이트'],
        ['12개월', '5,000명', '8%', '400명', '1,960,000원', '+ 출판사 B2B 리포트'],
        ['18개월', '10,000명', '10%', '1,000명', '4,900,000원', '+ 인용문 API'],
    ],
    col_widths=[1.8, 2, 1.5, 1.5, 3, 5.7]
)

add_divider()

add_h2('⑩ 멘토링 활용 계획')
add_table(
    ['멘토링 항목', '기대 내용'],
    [
        ['Claude API + RAG 비용 최적화', '호출 최소화·캐싱·pgvector 전환 설계'],
        ['인용 서비스 저작권 범위', '공정이용 기준, 출판사 협력 가능 범위 검토'],
        ['출판사 B2B 계약 구조', '밑줄 데이터 리포트 상품화, 계약 조건 설계'],
        ['프리미엄 전환율 전략', '월 10개 제한 적정성, 전환 유도 UX 설계'],
    ],
    col_widths=[5, 10.5]
)

add_h2('⑪ 사업화 자금 활용')
add_table(
    ['항목', '금액', '비중', '상세 내용'],
    [
        ['AI 추천 엔진 + RAG 개발', '1,600만원', '40%', 'Phase 1~3 구현 (자동 분석·RAG·추천 엔진)'],
        ['AI 활용모델 운영 인프라', '800만원', '20%', 'AI 비즈니스 모델 구현을 위한 API 서버·pgvector DB·CDN 인프라'],
        ['마케팅·홍보', '800만원', '20%', '저자 파트너십(200) + 북스타그래머(250) + 북클럽(150) + 콘텐츠·광고(200)'],
        ['UI/UX 고도화', '800만원', '20%', '모바일 UX 개선, 프리미엄 기능 구현'],
        ['합계', '4,000만원', '100%', '정부 3,200만원 + 자부담 800만원'],
    ],
    col_widths=[4, 2, 1.5, 8]
)

add_h2('⑫ 성과 목표 및 향후계획')
add_table(
    ['지표', '현재', '6개월 목표', '12개월 목표'],
    [
        ['MAU', '~50명', '1,000명', '5,000명'],
        ['누적 밑줄 수', '6,227개', '30,000개', '150,000개'],
        ['유료 구독자', '0명', '50명', '400명'],
        ['월 수익', '0원', '245,000원', '1,960,000원'],
        ['공유 카드 생성', '—', '5,000장', '30,000장'],
    ],
    col_widths=[4, 3, 3.5, 4],
    sub_header_rows={2, 3}
)
add_table(
    ['단계', '시점', '내용'],
    [
        ['수익화 시작', '사업 완료 직후', '프리미엄 구독 + 어필리에이트 출시'],
        ['마케팅 인력 채용', '완료 후 3개월', '저자 파트너십·북클럽 전담'],
        ['출판사 B2B', '완료 후 12개월', '밑줄 데이터 + AI 리포트 상품화'],
        ['투자 유치 준비', '완료 후 18개월', 'MAU 10,000명·유료 1,000명 달성 후'],
    ],
    col_widths=[3.5, 3.5, 8.5]
)

add_divider()
add_note('마감: 2026년 7월 3일(금) 16시  ｜  신청: www.sbiz24.kr  ｜  사전 필수: edu.sbiz.or.kr AI학습관 강의 1개 수강 완료')

# ─── 저장 ───────────────────────────────────────
out_path = r'd:\dev\underline\grant-application-2026-v2.docx'
doc.save(out_path)
print(f'저장 완료: {out_path}')
